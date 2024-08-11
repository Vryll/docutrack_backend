from ninja import Router, Form, File, Query
from .models import *
from .forms import *
from .schema import *
import io
from django.core.files import File as django_file
from user_accounts.models import Admin_Data
from ninja.files import UploadedFile
from django.core.exceptions import SuspiciousFileOperation
from django.core.files.base import ContentFile
import base64
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.http import Http404
from django.core import serializers
import json
from user_accounts.models import *
from django.http import JsonResponse
from json.decoder import JSONDecodeError
from django.shortcuts import get_object_or_404
from uuid import uuid4
import uuid
import time
from docx import section, Document as DocxDocument
from docx.shared import Inches
from io import BytesIO
import requests
from datetime import datetime
import logging
from docx2pdf import convert
from django.core.files.base import ContentFile
from django.db.models import Q
import pythoncom



router = Router()

@router.get('/getDocuStatus', response=List[DocumentStatus_Schema])
def List_Status(request):
    return Document_Status.objects.all()

@router.get('/getDocuTypes', response=List[DocumentTypes_Schema])
def List_Types(request):
    return Document_Type.objects.all()

@router.get('/getDocumentDetails', response=List[DocumentDetails_Schema])
def List_Details(request):
    return Document_Details.objects.all()

@router.get('/getDocumentSender', response=List[DocumentSender_Schema])
def List_DocuSender(request):
    docu_sender = Admin_Data.objects.all()
    return [{"office_name": obj.office_name.office_list} for obj in docu_sender]

#FETCHING DATA IN A SPECIFIC INSTANCE
@router.get('/getDocumentent_Details/{docuId}')
def getDocumentDetailsInstance(request, docuId: int):
    print(docuId)
    try:
        document_details = Document_Details.objects.get(id=docuId)

        print("docu_id:", document_details.id)
        print("docu_title:", document_details.docu_title)
        
       
        docu_type = None
        docu_status = None
        if document_details.status:
            docu_type = document_details.type.docu_type
            docu_status = document_details.status.docu_status
        
        print("docu_type:", docu_type)
        print("docu_status", docu_status)

        print("docu_sender:", document_details.docu_sender)
        print("docu_dateNtime_released:", document_details.docu_dateNtime_released)
        
        
        docu_recipient_list = [recipient.office_list for recipient in document_details.docu_recipient.all()]
        
        print("docu_recipient:", docu_recipient_list)
        
        print("docu_file:", document_details.docu_file.url)

        return {
            "docu_id": document_details.id,
            "memorandum_number": document_details.memorandum_number,
            "docu_title": document_details.docu_title,
            "docu_type": docu_type,
            "docu_status": docu_status,
            "docu_sender": document_details.docu_sender,
            "docu_dateNtime_released": document_details.docu_dateNtime_released,
            "docu_dateNtime_created":document_details.docu_dateNtime_created,
            "docu_recipient": docu_recipient_list,
            "docu_file": document_details.docu_file.url,
        }
    except Document_Details.DoesNotExist:
        raise Http404("Document_Details does not exist")


    
#FETCHING REQUESTED DOCUMENT
@router.get('/getRequestedDocument/req_docuId={reqDocuId}/{process_type}')
def GetRequestedDocument(request, reqDocuId: int, process_type: str):
    try:
       
        requested_document_details = Requested_Document.objects.get(
            Q(id=reqDocuId) & Q(requested_process_type__process_type=process_type)
        )


        requested_document_data = {
            "req_docu_id": requested_document_details.id,
            "docu_request_topic": requested_document_details.docu_request_topic,
            "requested": requested_document_details.requested,
            "docu_request_recipient": requested_document_details.docu_request_recipient,
            "docu_request_deadline": requested_document_details.docu_request_deadline,
            "docu_request_comment": requested_document_details.docu_request_comment,
            "docu_request_file": requested_document_details.docu_request_file.url,
            "status": requested_document_details.status.req_docu_status,
            "process_type": requested_document_details.requested_process_type.process_type,
        }

        return {"requested_document_data": requested_document_data}

    except Requested_Document.DoesNotExist:
        pass  

    try:
      
        forwarded_request_document_details = Forward_Request_Document.objects.get(
            Q(id=reqDocuId) & Q(forwarded_process_type__process_type=process_type)
        )

        forwarded_document_data = {
            "req_docu_id": forwarded_request_document_details.id,
            "docu_request_topic": forwarded_request_document_details.forwarded_subject,
            "requested": forwarded_request_document_details.forwarded_requested,
            "docu_request_recipient": forwarded_request_document_details.forwarded_docu_request_recipient,
            "docu_request_deadline": forwarded_request_document_details.forwarded_date_requested,
            "docu_request_comment": "", 
            "docu_request_file": forwarded_request_document_details.forwarded_requested_docu_file.url,
            "status": forwarded_request_document_details.forwarded_requested_docu_status.req_docu_status,
            "process_type": forwarded_request_document_details.forwarded_process_type.process_type,
        }

        return {"forwarded_document_data": forwarded_document_data}

    except Forward_Request_Document.DoesNotExist:
        pass  

   
    return {"error": "No matching data found for reqDocuId={} and process_type={}".format(reqDocuId, process_type)}


    

#ENDPOINT FOR FETCHING SOFT DELETED OUTGOING DOCUMENT DETAILS DATA
@router.get('/getSoftDeleted_outgoing_document_details/user_id={userId}')
def get_soft_deleted_outgoing_document_details(request, userId: int):
    admin_data_instance = Admin_Data.objects.get(user_id=userId)
    admin_office_name_id = admin_data_instance.office_name_id
    print(admin_office_name_id)

    office_name_instance = Office.objects.get(id=admin_office_name_id)
    office_name = office_name_instance.office_list
    print(office_name)

    archived_document_details_list = []
    archived_data = {"archived_docu_count": 0, "archived_document_details": []} 

   
    document_details_instances = Document_Details.deleted_objects.filter(
        Q(docu_sender=office_name) | Q(docu_recipient__id=admin_office_name_id), is_deleted=True
    ).distinct()
    archived_docu_count = document_details_instances.count()
    print(document_details_instances)

    for document_details_instance in document_details_instances:
        docu_recipient = [recipient.office_list for recipient in document_details_instance.docu_recipient.all()]

        result = {
            "docu_id": document_details_instance.id,
            "memorandum_number": document_details_instance.memorandum_number,
            "docu_title": document_details_instance.docu_title,
            "docu_type": document_details_instance.type.docu_type,
            "docu_status": document_details_instance.status.docu_status,
            "docu_sender": document_details_instance.docu_sender,
            "docu_dateNtime_released": document_details_instance.docu_dateNtime_released,
            "docu_dateNtime_created": document_details_instance.docu_dateNtime_created,
            "docu_recipient": docu_recipient,
            "docu_file": document_details_instance.docu_file.url if document_details_instance.docu_file else None,
            "is_deleted": document_details_instance.is_deleted,
            "deleted_at": document_details_instance.deleted_at,
        }

        archived_document_details_list.append(result)


    archived_data["archived_docu_count"] = archived_docu_count
    archived_data["archived_document_details"] = archived_document_details_list

    return archived_data



#CREATING DOCUMENT
@router.post('/generate_qr_code')
def GenerateQRCode(request, form: CreateDocumentSchema = Form(...)):

    # if not docu_file.name.endswith('.pdf'):
    #         return {"error": "Only PDF files are allowed for upload."}
    try:
        filestr = form.docu_file
        file_data = base64.b64decode(filestr)
        file_buffer = io.BytesIO(file_data)

 
        file_data_obj = django_file(file_buffer, name = form.docu_file_name)

       

    except Exception as e:
        print(f"An error occurred: {e}")
   

    office_instance = Office.objects.get(office_list=form.docu_sender)
    
    recipient_office_instance = []

    for docu_recipient in form.docu_recipient:
        
        individual_recipients = [recipient.strip() for recipient in docu_recipient.split(',')]

        for individual_recipient in individual_recipients:
            try:
                office = Office.objects.get(office_list=individual_recipient)
                recipient_office_instance.append(office)
            except Office.DoesNotExist:
               
                pass
    print (recipient_office_instance)

   
    document_details = Document_Details.objects.create( 
        memorandum_number=form.memorandum_number,
        docu_title=form.docu_title,
        type=Document_Type.objects.get(docu_type=form.docu_type),
        status=Document_Status.objects.get(docu_status=form.docu_status),
        docu_dateNtime_released=form.docu_dateNtime_released,
        docu_sender=office_instance.office_list,
        docu_file=file_data_obj
    )

   
    document_details.docu_recipient.set(recipient_office_instance)

    print(document_details.docu_sender)
    print(document_details.docu_recipient)
    print(document_details.memorandum_number)

    document_details_data = {
        "docu_id": document_details.id,
        "memorandum_number": document_details.memorandum_number,
        "docu_title": document_details.docu_title,
        "type": document_details.type.docu_type,
        "status": document_details.status.docu_status,
        "docu_dateNtime_released": document_details.docu_dateNtime_released,
        "docu_sender": document_details.docu_sender,
        "docu_recipient": [recipient.office_list for recipient in document_details.docu_recipient.all()],
        "docu_file": document_details.docu_file.name
      
    }

    return {"success": "Document was successfully saved!", "Document_Details": document_details_data}



@router.post('/getJwtTokenNOfficeName')
def GetJWTToken(request):
    try:  
   
        data = json.loads(request.body)
        
    
        user_token = data.get("jwt_token")
        
        try:
           
            user = User.objects.get(jwt_token=user_token)

           
            user_id = user.id

            admin_data_instance = Admin_Data.objects.get(user_id=user_id)

            admin_office_name = admin_data_instance.office_name

            office_instance = Office.objects.get(office_list=admin_office_name)

            requested_office_name = office_instance.office_list
            
          
            return {"requested": requested_office_name}
            
        except User.DoesNotExist:
            return JsonResponse({"message": "User not found"}, status=404)
       
    except Exception as e:
        return JsonResponse({"message": str(e)}, status=400)
    
#ENDPOINT TO FETCH LIST OF REQUESTED DOCUMENT FILES
@router.get('/getList_requested_document_file/{reqDocuID}')
def GetListRequestedDocumentFile(request, reqDocuID: int):
    requested_document_instance = Requested_Document.objects.get(id=reqDocuID)
    request_docu_id = requested_document_instance.id

    print(request_docu_id)

    default_requested_docu_file = {
        "docu_request_id": requested_document_instance.id,
        "docu_request_topic": requested_document_instance.docu_request_topic,
        "requested": requested_document_instance.requested,
        "docu_request_recipient": requested_document_instance.docu_request_recipient,
        "docu_request_deadline": requested_document_instance.docu_request_deadline,
        "docu_request_file": requested_document_instance.docu_request_file.url
    }

    additional_request_docu_file_list = []
    
    forward_request_document_instances = Forward_Request_Document.objects.filter(requested_document_id=request_docu_id)

    if forward_request_document_instances:
        for forward_request_document_instance in forward_request_document_instances:
            additional_request_docu_file = {
                "forwarded_requested_docu_id": forward_request_document_instance.id,
                "forwarded_subject": forward_request_document_instance.forwarded_subject,
                "forwarded_requested": forward_request_document_instance.forwarded_requested,
                "forwarded_docu_request_recipient": forward_request_document_instance.forwarded_docu_request_recipient,
                "forwarded_date_requested": forward_request_document_instance.forwarded_date_requested,
                "forwarded_requested_docu_file": forward_request_document_instance.forwarded_requested_docu_file.url if forward_request_document_instance.forwarded_requested_docu_file else None,
            }
            additional_request_docu_file_list.append(additional_request_docu_file)

    return {"default_requested_docu_file": default_requested_docu_file, "additional_request_docu_file": additional_request_docu_file_list}


#ENDPOINT FOR FETCHING THE SELECTED FORWARD REQUESTED DOCUMENT FILE 
@router.get('getForwarded_requested_document/{forwarded_reqDocuID}')
def GetForwardedRequestedDocument(request, forwarded_reqDocuID:int):
    forwarded_requested_docu_instance = Forward_Request_Document.objects.get(id = forwarded_reqDocuID)

    forwarded_requested_docu_file = {
        "forwarded_requested_docu_file":forwarded_requested_docu_instance.forwarded_requested_docu_file.url
    }

    return {"forwarded_requested_docu_file": forwarded_requested_docu_file}

#ENDPOINT FOR FETCHING SELECTED DEFAULT REQUESTED DOCUMENT FILE
@router.get('getRequested_document/{reqDocuID}')
def GetRequestedDocument(request, reqDocuID:int):
    requested_document_instance = Requested_Document.objects.get(id = reqDocuID)

    requested_docu_file = {
        "docu_request_file":requested_document_instance.docu_request_file.url
    }

    return {"requested_docu_file": requested_docu_file}

#create a request for document
@router.post('/create_request_document')
def RequestDocument(request, form: CreateRequestDocumentSchema = Form(...)):
        
        filestr = form.docu_request_file
        file_data = base64.b64decode(filestr)
        file_buffer = io.BytesIO(file_data)

        file_data_obj = django_file(file_buffer, name = form.docu_request_file_name)


        Requested_Document.objects.create(
            status=Requested_Document_Status.objects.get(req_docu_status=form.status),
            requested_process_type=Requested_Docu_Process_type.objects.get(process_type=form.requested_process_type),
            docu_request_topic=form.docu_request_topic,
            requested=form.requested,
            docu_request_recipient=form.docu_request_recipient,
            docu_request_deadline=form.docu_request_deadline,
            docu_request_file=file_data_obj,
        )
        return {"success": "Requested Document was successfully saved!"}
  


#ENDPOINT TO CREATE FORWARDED REQUEST FOR REQUESTED DOCUMENT
@router.post('/create_forwarded_request_document/{reqDocuID}')
def CreateForwardedRequestDocument(request, reqDocuID:int ,form: CreateForwardedRequestDocumentSchema = Form(...)):
    requested_document_instance = Requested_Document.objects.get(id = reqDocuID)

    request_docu_status_instance = Requested_Document_Status.objects.get(req_docu_status = form.forwarded_requested_docu_status)

    filestr = form.forwarded_requested_docu_file
    file_data = base64.b64decode(filestr)
    file_buffer = io.BytesIO(file_data)

    file_data_obj = django_file(file_buffer, name = form.forwarded_requested_docu_file_name)

    Forward_Request_Document.objects.create(
        forwarded_requested_docu_status = request_docu_status_instance,
        requested_document = requested_document_instance,
        forwarded_subject = form.forwarded_subject,
        forwarded_requested = form.forwarded_requested,
        forwarded_docu_request_recipient = form.forwarded_docu_request_recipient,
        forwarded_date_requested = form.forwarded_date_requested,
        forwarded_requested_docu_file = file_data_obj,
        forwarded_process_type=Requested_Docu_Process_type.objects.get(process_type = form.forwarded_process_type),
    )
    return {"success": "forwarded request is successfully saved!"}



#UPDATE FOR QR CODE
@router.post('/update_generated_qr_code/{docu_id}')
def UpdateGenerateQrCode(request, docu_id: int, form: UpdateGenerationQrCodeSchema = Form(...)):

        document_details_instance = Document_Details.objects.get(id=docu_id)
        print(form)

        imgstr = form.docu_qr_code
        image_data = base64.b64decode(imgstr)
        image_buffer = io.BytesIO(image_data)
        unique_qr_name = f"qr_code_{uuid4().hex}.png"
        image_file = django_file(image_buffer, name=unique_qr_name)
        file_data = image_file

        document_details_instance.docu_qr_code = file_data
        document_details_instance.save()
        print(file_data.name)


        return {"success": "Document updated successfully"}


#ENDPOINT TO CREATE FORWARDED REQUESTED DOCUMENT RECORDS
@router.post('create_forwarded_requested_document_records/{reqDocuId}')
def CreateForwardedRequestedDocumentRecords(request, reqDocuId:int, form: CreateForwardedDocumentRecordSchema = Form(...)):
    requested_document_instance = Requested_Document.objects.get(id = reqDocuId)
    req_docuId = requested_document_instance.id
    print(req_docuId)

    forwarded_record_status_instance = Requested_Document_Status.objects.get(req_docu_status=form.forwarded_record_status)

    forwarded_requested_document_instance = Forward_Request_Document.objects.get(requested_document_id = req_docuId)

    Forwarded_Request_Document_Record.objects.create(
        requested_document = requested_document_instance,
        forward_request_document = forwarded_requested_document_instance,
        forwarded_record_status=forwarded_record_status_instance
    )
    return{"message": "forwarded requested document was successfully saved!"}



#UPDATE DATA ENDPOINTS
@router.post('/update_requested_document/{reqDocuId}')
def update_requested_document(request, reqDocuId: int):
    try:
        data = json.loads(request.body)
        status_value = data.get("status_value")

        request_document_instance = Requested_Document.objects.get(id=reqDocuId)

        try:
            req_docu_status = Requested_Document_Status.objects.get(req_docu_status=status_value)
        except Document_Status.DoesNotExist:
            return JsonResponse({"error": f"Document_Status with docu_status '{status_value}' does not exist."}, status=400)

        request_document_instance.status = req_docu_status
        request_document_instance.save()

        updated_status = request_document_instance.status.req_docu_status
        return JsonResponse({"success": True, "updated_status": updated_status})
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
    

#ENDPOINT FOR UPDATED FORWARDED REQUESTED DOCUMENT STATUS

@router.post('/word_file_processing/{docuId}')
def WordFileProcessing(request, docuId: int):
    print(docuId)
    pythoncom.CoInitialize()
    try:
        document_detail_instance = get_object_or_404(Document_Details, id=docuId)

      
        docx_file_url = document_detail_instance.docu_file.url
        qr_code_url = document_detail_instance.docu_qr_code.url

        print(docx_file_url)
        print(qr_code_url)

        if not docx_file_url or not qr_code_url:
            return JsonResponse({"error": "DOCX file URL or QR code URL not found"}, status=404)

       
        docx_response = requests.get(docx_file_url)
        qr_code_response = requests.get(qr_code_url)

        if docx_response.status_code != 200 or qr_code_response.status_code != 200:
            return JsonResponse({"error": "Failed to download files"}, status=500)

        docx_file_content = docx_response.content
        qr_code_content = qr_code_response.content

        docx = DocxDocument(BytesIO(docx_file_content))
        
        section = docx.sections[0]
        footer = section.footer

        
        qr_image_buffer = BytesIO(qr_code_content)
        qr_image_buffer.seek(0)
        
       
        content_width = docx.sections[0].page_width + docx.sections[0].left_margin + docx.sections[0].right_margin

 
        image_width = Inches(1)
        image_height = Inches(1)

     
        insert_qr = footer.paragraphs[0]
        run = insert_qr.add_run()
        run.add_picture(qr_image_buffer, width=image_width, height=image_height)


        
        unique_filename = str(uuid.uuid4()) + "_modified.docx"

       
        modified_docx_file = unique_filename
       

        docx.save(modified_docx_file)

        pdf_filename = f'pdf/{modified_docx_file}.pdf'

        convert(modified_docx_file,pdf_filename)
        
        try:
            with open(pdf_filename, 'rb') as file:
                
                content = file.read()
               
                document_detail_instance.modified_docu_file.save(f'{modified_docx_file}.pdf',ContentFile(content))
                
        except FileNotFoundError:
            print(f"File not found: {pdf_filename}")

        response_data = {
            "modified_docu_file": document_detail_instance.modified_docu_file.url,
        }


        return response_data
    except Document_Details.DoesNotExist as e:
        print(e)
        return JsonResponse({"error": "Document not found"}, status=404)
    except Exception as e:
        print(e)
        return JsonResponse({"error": str(e)}, status=500)
    finally:
        pythoncom.CoUninitialize()

#ENDPOINT FOR FETCHING DATE AND TIME FROM FLUTTER APP
@router.post('/DateTime/user_id={userId}')
def DateTime(request, userId:int):
    try:
        data = json.loads(request.body)
        dateTime = data.get("formattedDate")

        admin_data_instance = Admin_Data.objects.get(user_id=userId)
        admin_office_id = admin_data_instance.office_name_id
        print(admin_office_id)

        office_instance = Office.objects.get(id=admin_office_id)
        admin_office_name = office_instance.office_list
        print(admin_office_name)

        if dateTime is not None:
            date = datetime.strptime(dateTime, "%Y-%m-%d").date()
            print(dateTime)

           
            document_details_instances = Document_Details.objects.filter(docu_sender=admin_office_name, docu_dateNtime_created__date=date)

           
            count = document_details_instances.count()
            print(count)

            return {"count": count}
        else:
            return JsonResponse({"error": "dateTime is None"}, status=400)

    except Admin_Data.DoesNotExist:
        return JsonResponse({"error": "Admin_Data does not exist"}, status=404)
    except Office.DoesNotExist:
        return JsonResponse({"error": "Office does not exist"}, status=404)
    except Document_Details.DoesNotExist:
        return JsonResponse({"error": "Document_Details does not exist"}, status=404)
    except Exception as e:
        logging.exception("An unexpected error occurred: %s", str(e))
        return JsonResponse({"error": "Internal Server Error"}, status=500)


from django.db.models import Prefetch

@router.get('filterOutgoingDocuDetails/{userId}')
def filterOutgoingDocuDetails(request, userId: int, filters: OutgoingDocuDetailsFilterSchema = Query(...)):
    try:
        user_instance = User.objects.get(id=userId)
    except User.DoesNotExist:
        return JsonResponse({"error": "User not found"}, status=404)

    admin_data_instance = get_object_or_404(Admin_Data, user_id=user_instance.id)
    admin_office_name = admin_data_instance.office_name

    try:
        office_instance = Office.objects.get(office_list=admin_office_name)
    except Office.DoesNotExist:
        return JsonResponse({"error": "No matching office data found"}, status=404)

    office_name = office_instance.office_list

    
    document_details_instances = Document_Details.objects.filter(docu_sender=office_name)

    if filters.memorandum_number:
        document_details_instances = document_details_instances.filter(memorandum_number=filters.memorandum_number)

    if filters.type:
        if isinstance(filters.type, str):
            document_details_instances = document_details_instances.filter(type__docu_type=filters.type)
        else:
            document_details_instances = document_details_instances.filter(type__docu_type=filters.type.docu_type)

    if filters.docu_recipient:
        if isinstance(filters.docu_recipient, str):
            document_details_instances = document_details_instances.filter(docu_recipient__office_list=filters.docu_recipient)
        else:
            document_details_instances = document_details_instances.filter(docu_recipient__office_list=filters.docu_recipient.office_list)


    
    document_details_instances = document_details_instances.prefetch_related(
        Prefetch('type', queryset=Document_Type.objects.all()),
        Prefetch('docu_recipient', queryset=Office.objects.all()),
    )

   
    filtered_docu_detail = document_details_instances.values(
        'docu_recipient__office_list',  
        'type__docu_type', 
        'memorandum_number',
        'docu_title',
        'docu_dateNtime_created',
       
    ).first()  

    if filtered_docu_detail is not None:
      
        result_data = {
            "docu_recipient": filtered_docu_detail.get('docu_recipient__office_list'),
            "type": filtered_docu_detail.get('type__docu_type'),
            "memorandum_number": filtered_docu_detail.get('memorandum_number'),
            "docu_title": filtered_docu_detail.get('docu_title'),
            "docu_dateNtime_created": filtered_docu_detail.get('docu_dateNtime_created'),
       
        }
        return {"filtered_data": result_data}
    else:
        return {"error": "No matching data found"}

#ENDPOINTS FOR DELETION OF DATA FOR DOCUMENT DETAILS 
@router.delete("deleting_document_details/documentID={docuId}")
def DeleteDocuDetails(request, docuId: int):
    docu_details_instance = get_object_or_404(Document_Details, id=docuId)

    if docu_details_instance is not None:
        docu_details_instance.delete()
        return {'message': f'The document detail with ID {docuId} was successfully deleted'}
    else:
        return {'message': f'The document detail with ID {docuId} does not exist'}
    
#ENDPOINTS FOR RESTORING OF DATA FOR DOCUMENT DETAILS 
@router.get("restoring_document_details/documentID={docuId}")
def DeleteDocuDetails(request, docuId: int):
    docu_details_instance = Document_Details.deleted_objects.get(id=docuId)

    if docu_details_instance is not None:
        docu_details_instance.restore()
        return {'message': f'The document detail with ID {docuId} was successfully restored'}
    else:
        return {'message': f'The document detail with ID {docuId} does not exist'}
    

#ENDPOINTS FOR HARD DELETING OF DOCUMENT DETAILS
@router.delete("/hard_delete_document_details/documentID={docuId}")
def HardDeleteDocumentDetails(request, docuId):
    docu_details_instance = Document_Details.deleted_objects.get(id = docuId)

    if docu_details_instance is not None:
        docu_details_instance.hard_delete()
        return {'message': f'The document detail with ID {docuId} was successfully deleted'}
    else:
        return {'message': f'The document detail with ID {docuId} does not exist'}

#ENDPOINTS FOR DELETION OF DATA FOR REQUESTED DOCUMENT DETAILS 
@router.delete("deleting_requested_document_details/requested_documentID={req_docuId}")
def DeleteDocuDetails(request, req_docuId: int):
    docu_details_instance = get_object_or_404(Requested_Document, id=req_docuId)

    if docu_details_instance is not None:
        docu_details_instance.delete()
        return {'message': f'The document detail with ID {req_docuId} was successfully deleted'}
    else:
        return {'message': f'The document detail with ID {req_docuId} does not exist'}
    
#ENDPOINTS FOR RESTORING OF DATA FOR REQUESTED DOCUMENT DETAILS 
@router.get("restoring_requested_document_details/requested_documentID={req_docuId}")
def DeleteDocuDetails(request, req_docuId: int):
    docu_details_instance = Requested_Document.deleted_objects.get(id=req_docuId)

    if docu_details_instance is not None:
        docu_details_instance.restore()
        return {'message': f'The document detail with ID {req_docuId} was successfully restored'}
    else:
        return {'message': f'The document detail with ID {req_docuId} does not exist'}
    
#ENDPOINTS FOR HARD DELETE OF DATA FOR REQUESTED DOCUMENT DETAILS 
@router.get("hard_delete_requested_document_details/requested_documentID={req_docuId}")
def DeleteDocuDetails(request, req_docuId: int):
    docu_details_instance = Requested_Document.deleted_objects.get(id=req_docuId)

    if docu_details_instance is not None:
        docu_details_instance.hard_delete()
        return {'message': f'The document detail with ID {req_docuId} was successfully restored'}
    else:
        return {'message': f'The document detail with ID {req_docuId} does not exist'}
    


#ENDPOINT FOR UPDATING THE STATUS FOR REQUESTED DOCUMENT IN FORWARDED STATUS
@router.post('update_forward_status_reqDocu/{reqDocuID}')
def UpdateReceivedStatusReqDocu(request,reqDocuID:int ):
   

    data = json.loads(request.body)

    status_value = data.get("status_value")

    requested_document_status_instance = Requested_Document_Status.objects.get(req_docu_status = status_value)

    requested_document_instance = Requested_Document.objects.get(id = reqDocuID)
    requested_document_instance.status = requested_document_status_instance
    requested_document_instance.save()

    return {'message':'Received Docu status has successfully updated to received'}



#ENDPOINT FOR UPDATING THE STATUS FOR FORWARDED REQUESTED DOCUMENT
# @router.post('update_forwarded_request_status/{reqDocuID}')
# def UpdateForwardedRequestStatus(request, reqDocuID:int):
#      status_value = "Forwarded"
     
#      requested_document_status_instance = Requested_Document_Status.objects.get(req_docu_status = status_value)
     
#      requested_document_instance = Forward_Request_Document.objects.filter(requested_document_id = reqDocuID).first('forwarded_request_time_created')
#      requested_document_instance.status = requested_document_status_instance
#      requested_document_instance.save()

#      return {'message':'ForwardedReceived Docu status has successfully updated to received'}

#ENDPOINT FOR UPDATING THE FORWARDED REQUESTED DOCUMENT
@router.post('update_forwarded_request_document/{reqDocuID}')
def UpdateForwardedRequestDocument(request, reqDocuID: int):
    
    data = json.loads(request.body)

    status_value = data.get("status_value")
    print(status_value)

    requested_document_status_instance = Requested_Document_Status.objects.get(req_docu_status = status_value)

    forwarded_document_record_instance = Forwarded_Request_Document_Record.objects.filter(requested_document_id = reqDocuID).latest('forwarded_request_time_created')
    forwarded_record_id = forwarded_document_record_instance.forward_request_document_id
    print(forwarded_record_id)

    forward_request_document_instance = Forward_Request_Document.objects.get(id = forwarded_record_id)
    forward_request_document_instance.forwarded_requested_docu_status = requested_document_status_instance
    forward_request_document_instance.save()

    return{'message': 'forwarded request status is successfully saved to FORWARDED'}

#ENDPOINT FOR UPDATING THE RECEIVED STATUS FOR REQUESTED DOCUMENT
@router.post('update_forwarded_request_receive_status/{reqDocuID}')
def UpdateForwardedRequestReceiveStatus(request, reqDocuID: int):
    # data = json.loads(request.body)

    # status_value = data.get("status_value")
    status_value = "Received"
    print(status_value)

    requested_document_status_instance = get_object_or_404(Requested_Document_Status, req_docu_status=status_value)
    requested_document_instance = get_object_or_404(Requested_Document, id=reqDocuID)
    
    requested_document_instance.status = requested_document_status_instance
    requested_document_instance.save()

    forward_request_document_instances = Forward_Request_Document.objects.filter(requested_document=reqDocuID)
    for forward_request_document_instance in forward_request_document_instances:
        forward_request_document_instance.forwarded_requested_docu_status = requested_document_status_instance
        forward_request_document_instance.save()

    return {"message": "Requested Document was successfully updated to RECEIVED"}

#ENDPOINT FOR UPDATING THE RETURNED STATUS FOR REQUESTED DOCUMENT
@router.post('update_forwarded_request_returned_status/{reqDocuID}')
def UpdateForwardedRequestReceiveStatus(request, reqDocuID: int):
    # data = json.loads(request.body)

    # status_value = data.get("status_value")
    status_value = "Returned"
    print(status_value)

    requested_document_status_instance = get_object_or_404(Requested_Document_Status, req_docu_status=status_value)
    requested_document_instance = get_object_or_404(Requested_Document, id=reqDocuID)
    
    requested_document_instance.status = requested_document_status_instance
    requested_document_instance.save()

    forward_request_document_instances = Forward_Request_Document.objects.filter(requested_document=reqDocuID)
    for forward_request_document_instance in forward_request_document_instances:
        forward_request_document_instance.forwarded_requested_docu_status = requested_document_status_instance
        forward_request_document_instance.save()

    return {"message": "Requested Document was successfully updated to RECEIVED"}

#ENDPOINT TO CREATE COMMENT FOR RETURNED DOCUMENTS
@router.post('create_returned_comment/{reqDocuID}')
def CreateReturnedComment(request, reqDocuID: int, form: CreateReturnedCommentSchema = Form(...)):
    request_document_instance = Requested_Document.objects.get(id=reqDocuID)

    create_data = {
        'docu_request_comment': form.docu_request_comment,
    }

    
    request_document_instance.docu_request_comment = form.docu_request_comment
    request_document_instance.save()

    return {'message': 'Comment created successfully'}


#ENDPOINT FOR FORWARDED DOCUMENT RECORDS 
@router.get('getForwarded_document_records/{reqDocuID}')
def GetForwardedDocumentRecords(request, reqDocuID: int):
    try:
        request_document_instance = Requested_Document.objects.get(id=reqDocuID)
        requested_docuId = request_document_instance.id

        forward_request_instance_record_instance = Forwarded_Request_Document_Record.objects.filter(requested_document=requested_docuId)
        
        forward_request_list = []
        
        for record_instance in forward_request_instance_record_instance:
            forward_request_docu_instance = record_instance.forward_request_document
            forward_request_data = {
                "forwarded_requested_docu_status": forward_request_docu_instance.forwarded_requested_docu_status.req_docu_status,
                "forwarded_subject": forward_request_docu_instance.forwarded_subject,
                "forwarded_requested": forward_request_docu_instance.forwarded_requested,
                "forwarded_docu_request_recipient": forward_request_docu_instance.forwarded_docu_request_recipient,
                "forwarded_date_requested": forward_request_docu_instance.forwarded_date_requested,
                "forwarded_requested_docu_file": forward_request_docu_instance.forwarded_requested_docu_file.url,
                "forwarded_requested_docu_time_Stamp": forward_request_docu_instance.forwarded_requested_docu_time_Stamp
            }
            forward_request_list.append(forward_request_data)

        return JsonResponse({"forward_requested_data": forward_request_list})
    except Requested_Document.DoesNotExist:
        return JsonResponse({"error": "Requested Document not found"}, status=404)
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
